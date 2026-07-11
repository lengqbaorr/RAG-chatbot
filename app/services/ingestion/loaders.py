from __future__ import annotations

import hashlib
import mimetypes
from abc import ABC, abstractmethod
from io import BytesIO
from pathlib import Path
from urllib.parse import urlparse

import requests
import trafilatura
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from PIL import Image, UnidentifiedImageError
from pydantic import AnyUrl, BaseModel, Field
from pypdf import PdfReader

from app.core.config import settings
from app.schemas.document import Document, DocumentMetadata, DocumentType
from app.services.ingestion.exceptions import (
    DocumentLoadError,
    OCRNotAvailableError,
    UnsupportedDocumentTypeError,
)


TEXT_EXTENSIONS = {".txt": DocumentType.txt, ".md": DocumentType.markdown}
IMAGE_EXTENSIONS = {
    ".bmp",
    ".gif",
    ".jpeg",
    ".jpg",
    ".png",
    ".tif",
    ".tiff",
    ".webp",
}


class LoaderInput(BaseModel):
    source: str
    document_type: DocumentType | None = None
    title: str | None = None
    user_id: str | None = None
    ocr_languages: str | None = Field(
        default=None,
        description="Tesseract language codes, for example 'eng' or 'eng+vie'.",
    )


class BaseDocumentLoader(ABC):
    @abstractmethod
    def load(self, loader_input: LoaderInput) -> list[Document]:
        raise NotImplementedError


class TextFileLoader(BaseDocumentLoader):
    def load(self, loader_input: LoaderInput) -> list[Document]:
        path = _require_existing_file(loader_input.source)
        try:
            text = path.read_text(encoding="utf-8-sig")
        except UnicodeDecodeError:
            text = path.read_text(encoding="utf-8", errors="replace")
        except OSError as exc:
            raise DocumentLoadError(f"Cannot read text file: {path}") from exc

        document_type = loader_input.document_type or _type_from_extension(path)
        return [
            _build_document(
                text=text,
                document_type=document_type,
                source=str(path),
                title=loader_input.title or path.name,
                user_id=loader_input.user_id,
                mime_type=mimetypes.guess_type(path.name)[0],
            )
        ]


class PDFLoader(BaseDocumentLoader):
    def load(self, loader_input: LoaderInput) -> list[Document]:
        path = _require_existing_file(loader_input.source)
        documents: list[Document] = []

        try:
            reader = PdfReader(str(path))
            documents = _documents_from_pdf_reader(
                reader,
                source=str(path),
                title=loader_input.title or path.name,
                user_id=loader_input.user_id,
            )
        except Exception as exc:
            raise DocumentLoadError(f"Cannot read PDF file: {path}") from exc

        return documents


class DocxLoader(BaseDocumentLoader):
    def load(self, loader_input: LoaderInput) -> list[Document]:
        path = _require_existing_file(loader_input.source)

        try:
            doc = DocxDocument(str(path))
        except Exception as exc:
            raise DocumentLoadError(f"Cannot read DOCX file: {path}") from exc

        parts: list[str] = []
        parts.extend(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return [
            _build_document(
                text="\n".join(parts),
                document_type=DocumentType.docx,
                source=str(path),
                title=loader_input.title or path.name,
                user_id=loader_input.user_id,
                mime_type=(
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
            )
        ]


class HTMLUrlLoader(BaseDocumentLoader):
    def load(self, loader_input: LoaderInput) -> list[Document]:
        url = _require_http_url(loader_input.source)
        headers = {"User-Agent": settings.http_user_agent}

        try:
            response = requests.get(url, headers=headers, timeout=settings.http_timeout_seconds)
            response.raise_for_status()
        except requests.RequestException as exc:
            raise DocumentLoadError(f"Cannot fetch URL: {url}") from exc

        content = getattr(response, "content", None)
        if content is None:
            content = response.text.encode("utf-8")
        response_url = getattr(response, "url", url)
        content_type = response.headers.get("content-type", "")
        if _is_pdf_response(content, content_type=content_type, response_url=response_url):
            try:
                reader = PdfReader(BytesIO(content))
                return _documents_from_pdf_reader(
                    reader,
                    source=url,
                    title=loader_input.title or _url_filename(response_url, default="document.pdf"),
                    user_id=loader_input.user_id,
                )
            except Exception as exc:
                raise DocumentLoadError(f"Cannot read PDF from URL: {url}") from exc

        html = response.text
        soup = BeautifulSoup(html, "lxml")
        extracted_text = trafilatura.extract(html, url=url) or _fallback_html_text(soup)
        title = loader_input.title or _html_title(soup) or url

        return [
            _build_document(
                text=extracted_text,
                document_type=DocumentType.url,
                source=url,
                title=title,
                user_id=loader_input.user_id,
                mime_type=response.headers.get("content-type"),
            )
        ]


class ImageOCRLoader(BaseDocumentLoader):
    def load(self, loader_input: LoaderInput) -> list[Document]:
        path = _require_existing_file(loader_input.source)
        _configure_tesseract()

        try:
            import pytesseract

            with Image.open(path) as image:
                text = pytesseract.image_to_string(
                    image,
                    lang=loader_input.ocr_languages or settings.ocr_languages,
                )
        except UnidentifiedImageError as exc:
            raise DocumentLoadError(f"Cannot identify image file: {path}") from exc
        except pytesseract.TesseractNotFoundError as exc:
            raise OCRNotAvailableError(
                "Tesseract executable was not found. Install Tesseract OCR and set "
                "TESSERACT_CMD in .env if it is not available on PATH."
            ) from exc
        except pytesseract.TesseractError as exc:
            raise DocumentLoadError(f"Tesseract OCR failed for image: {path}") from exc
        except OSError as exc:
            raise DocumentLoadError(f"Cannot read image file: {path}") from exc

        return [
            _build_document(
                text=text,
                document_type=DocumentType.image,
                source=str(path),
                title=loader_input.title or path.name,
                user_id=loader_input.user_id,
                mime_type=mimetypes.guess_type(path.name)[0],
            )
        ]


class DocumentLoaderService:
    def __init__(self) -> None:
        self._loaders: dict[DocumentType, BaseDocumentLoader] = {
            DocumentType.txt: TextFileLoader(),
            DocumentType.markdown: TextFileLoader(),
            DocumentType.pdf: PDFLoader(),
            DocumentType.docx: DocxLoader(),
            DocumentType.url: HTMLUrlLoader(),
            DocumentType.image: ImageOCRLoader(),
        }

    def load(self, loader_input: LoaderInput) -> list[Document]:
        document_type = loader_input.document_type or infer_document_type(loader_input.source)
        loader = self._loaders.get(document_type)
        if loader is None:
            raise UnsupportedDocumentTypeError(f"Unsupported document type: {document_type}")

        return loader.load(loader_input.model_copy(update={"document_type": document_type}))


def infer_document_type(source: str) -> DocumentType:
    if _is_http_url(source):
        return DocumentType.url

    path = Path(source)
    return _type_from_extension(path)


def _build_document(
    *,
    text: str,
    document_type: DocumentType,
    source: str,
    title: str | None = None,
    page_number: int | None = None,
    user_id: str | None = None,
    mime_type: str | None = None,
) -> Document:
    clean_text = _normalize_text(text)
    document_id = _document_id(
        source=source,
        document_type=document_type,
        page_number=page_number,
        text=clean_text,
        user_id=user_id,
    )

    return Document(
        text=clean_text,
        metadata=DocumentMetadata(
            document_id=document_id,
            document_type=document_type,
            source=source,
            title=title,
            page_number=page_number,
            user_id=user_id,
            mime_type=mime_type,
        ),
    )


def _document_id(
    *,
    source: str,
    document_type: DocumentType,
    page_number: int | None,
    text: str,
    user_id: str | None,
) -> str:
    raw = f"{user_id or ''}|{document_type}|{source}|{page_number or ''}|{text[:2048]}"
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def _normalize_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]
    return "\n".join(line for line in lines if line)


def _require_existing_file(source: str) -> Path:
    path = Path(source)
    if not path.exists() or not path.is_file():
        raise DocumentLoadError(f"File does not exist: {source}")
    return path


def _require_http_url(source: str) -> str:
    if not _is_http_url(source):
        raise DocumentLoadError(f"URL must start with http:// or https://: {source}")
    AnyUrl(source)
    return source


def _is_http_url(source: str) -> bool:
    parsed = urlparse(source)
    return parsed.scheme in {"http", "https"} and bool(parsed.netloc)


def _type_from_extension(path: Path) -> DocumentType:
    suffix = path.suffix.lower()
    if suffix in TEXT_EXTENSIONS:
        return TEXT_EXTENSIONS[suffix]
    if suffix == ".pdf":
        return DocumentType.pdf
    if suffix == ".docx":
        return DocumentType.docx
    if suffix in IMAGE_EXTENSIONS:
        return DocumentType.image

    raise UnsupportedDocumentTypeError(f"Unsupported file extension: {suffix or '<none>'}")


def _fallback_html_text(soup: BeautifulSoup) -> str:
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return soup.get_text(separator="\n", strip=True)


def _html_title(soup: BeautifulSoup) -> str | None:
    if soup.title and soup.title.string:
        return soup.title.string.strip()
    heading = soup.find("h1")
    return heading.get_text(strip=True) if heading else None


def _documents_from_pdf_reader(
    reader: PdfReader,
    *,
    source: str,
    title: str,
    user_id: str | None,
) -> list[Document]:
    documents: list[Document] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if not text.strip():
            continue
        documents.append(
            _build_document(
                text=text,
                document_type=DocumentType.pdf,
                source=source,
                title=title,
                page_number=index,
                user_id=user_id,
                mime_type="application/pdf",
            )
        )
    return documents


def _is_pdf_response(content: bytes, *, content_type: str, response_url: str) -> bool:
    media_type = content_type.split(";", maxsplit=1)[0].strip().lower()
    return (
        media_type == "application/pdf"
        or content.lstrip().startswith(b"%PDF-")
        or urlparse(response_url).path.lower().endswith(".pdf")
    )


def _url_filename(url: str, *, default: str) -> str:
    filename = Path(urlparse(url).path).name
    if not filename:
        return default
    return filename if filename.lower().endswith(".pdf") else f"{filename}.pdf"


def _configure_tesseract() -> None:
    if not settings.tesseract_cmd:
        return

    try:
        import pytesseract

        pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd
    except ImportError as exc:
        raise OCRNotAvailableError(
            "pytesseract is not installed. Install project dependencies first."
        ) from exc
