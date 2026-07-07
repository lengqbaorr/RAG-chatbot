from pathlib import Path

import fitz
import pytest
from docx import Document as DocxDocument

from app.schemas.document import DocumentType
from app.services.ingestion import (
    DocumentLoaderService,
    LoaderInput,
    UnsupportedDocumentTypeError,
    infer_document_type,
)


def test_load_txt_file(tmp_path: Path) -> None:
    source = tmp_path / "note.txt"
    source.write_text("  Hello RAG  \n\nProduction ready", encoding="utf-8")

    documents = DocumentLoaderService().load(LoaderInput(source=str(source)))

    assert len(documents) == 1
    assert documents[0].text == "Hello RAG\nProduction ready"
    assert documents[0].metadata.document_type == DocumentType.txt
    assert documents[0].metadata.title == "note.txt"


def test_load_markdown_file(tmp_path: Path) -> None:
    source = tmp_path / "note.md"
    source.write_text("# Title\n\nMarkdown content", encoding="utf-8")

    documents = DocumentLoaderService().load(LoaderInput(source=str(source)))

    assert len(documents) == 1
    assert documents[0].text == "# Title\nMarkdown content"
    assert documents[0].metadata.document_type == DocumentType.markdown


def test_load_docx_file(tmp_path: Path) -> None:
    source = tmp_path / "document.docx"
    doc = DocxDocument()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Key"
    table.rows[0].cells[1].text = "Value"
    doc.save(source)

    documents = DocumentLoaderService().load(LoaderInput(source=str(source)))

    assert len(documents) == 1
    assert "First paragraph" in documents[0].text
    assert "Second paragraph" in documents[0].text
    assert "Key | Value" in documents[0].text
    assert documents[0].metadata.document_type == DocumentType.docx


def test_load_pdf_file(tmp_path: Path) -> None:
    source = tmp_path / "document.pdf"
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "PDF content")
    pdf.save(source)
    pdf.close()

    documents = DocumentLoaderService().load(LoaderInput(source=str(source)))

    assert len(documents) == 1
    assert "PDF content" in documents[0].text
    assert documents[0].metadata.document_type == DocumentType.pdf
    assert documents[0].metadata.page_number == 1


def test_load_html_url(monkeypatch: pytest.MonkeyPatch) -> None:
    class FakeResponse:
        text = "<html><head><title>Demo</title></head><body><h1>Hello</h1><p>URL content</p></body></html>"
        headers = {"content-type": "text/html; charset=utf-8"}

        def raise_for_status(self) -> None:
            return None

    def fake_get(*args: object, **kwargs: object) -> FakeResponse:
        return FakeResponse()

    monkeypatch.setattr("app.services.ingestion.loaders.requests.get", fake_get)

    documents = DocumentLoaderService().load(LoaderInput(source="https://example.com/demo"))

    assert len(documents) == 1
    assert "URL content" in documents[0].text
    assert documents[0].metadata.document_type == DocumentType.url
    assert documents[0].metadata.title == "Demo"


def test_infer_document_type_for_url() -> None:
    assert infer_document_type("https://example.com/page") == DocumentType.url


def test_infer_document_type_for_image(tmp_path: Path) -> None:
    source = tmp_path / "scan.png"
    source.write_bytes(b"not actually loaded here")

    assert infer_document_type(str(source)) == DocumentType.image


def test_unsupported_extension(tmp_path: Path) -> None:
    source = tmp_path / "archive.zip"
    source.write_text("not supported", encoding="utf-8")

    with pytest.raises(UnsupportedDocumentTypeError):
        DocumentLoaderService().load(LoaderInput(source=str(source)))
